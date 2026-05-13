#!/usr/bin/env python3
"""Build the ADALM-Pluto dashboard plot guide DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
IMG_DIR = DOCS / "plot_guide_screenshots"
OUT = DOCS / "ADALM_Pluto_Dashboard_Plot_Guide.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
BLACK = RGBColor(0, 0, 0)
MUTED = RGBColor(85, 85, 85)
TABLE_FILL = "E8EEF5"
BORDER = "DADCE0"

MODES = [
    ("BPSK", "bpsk", "1 bit/symbol", "Two possible symbol decisions, mainly separated along the I axis."),
    ("QPSK", "qpsk", "2 bits/symbol", "Four possible phase states, one in each I/Q quadrant."),
    ("16-QAM", "16_qam", "4 bits/symbol", "Sixteen amplitude-and-phase states arranged as a grid in I/Q."),
]

PLOTS = [
    (
        "Power Spectrum / PSD",
        "spectrum",
        "This plot shows where the received signal power sits in frequency. It is the same idea as the "
        "power spectral density and spectrum plots from class: the x-axis is frequency in hertz and the y-axis "
        "is power in dB.",
        [
            "Look for the raised center band: that is the transmitted signal energy around baseband.",
            "The blue trace is the raw received Pluto samples; the red trace includes the software impairments selected in the dashboard.",
            "If frequency offset is added, the signal energy shifts left or right. If noise is added, the floor around the signal rises.",
        ],
        "A healthy display has one clear occupied band near the center and a lower surrounding noise floor.",
    ),
    (
        "I/Q Constellation",
        "constellation",
        "This plot shows the received symbols after the receiver has converted the signal into in-phase "
        "and quadrature components. The horizontal axis is I and the vertical axis is Q.",
        [
            "Yellow x marks show ideal decision locations; green dots show received symbol estimates from the Pluto signal.",
            "Tight clusters near the ideal points mean cleaner reception. Wider clouds mean more noise, timing error, or synchronization error.",
            "Rotation usually points to phase offset. A stretched or smeared pattern can come from frequency offset, filtering, or a noisy channel.",
        ],
        "The expected pattern changes with modulation: BPSK has two decisions, QPSK has four, and 16-QAM has a denser grid.",
    ),
    (
        "Eye Diagram",
        "eye",
        "This plot overlays many short windows of the matched-filter output. It is a class tool for seeing "
        "symbol timing, noise margin, and intersymbol interference.",
        [
            "The dashed yellow line marks the nominal sampling instant, one symbol period into the two-symbol window.",
            "A more open eye means the receiver has a safer time and amplitude margin for deciding symbols.",
            "A closed, fuzzy, or crossing-heavy eye means noise, timing error, bandwidth limitation, or ISI is making decisions harder.",
        ],
        "Use the eye diagram as a timing-quality view, not as a direct bit counter.",
    ),
    (
        "Throughput Counter",
        "throughput",
        "This plot is the live data-rate view. It keeps the useful counter you liked: yellow shows the "
        "instantaneous decoded rate in kb/s and purple shows the cumulative number of decoded bits.",
        [
            "The yellow line jumps because each Pluto receive/update cycle produces a different short-term rate.",
            "The purple line should generally climb over time while the dashboard is streaming.",
            "Higher bits per symbol can raise the possible bit rate, but the actual live value also depends on synchronization and buffer timing.",
        ],
        "This plot is the best place to answer: is the receiver still decoding data and how quickly is the bit count growing?",
    ),
]


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_para(doc, text="", style=None, before=0, after=6, line=1.25):
    p = doc.add_paragraph(style=style)
    set_paragraph_spacing(p, before=before, after=after, line=line)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=11, color=BLACK)
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    if level == 1:
        set_paragraph_spacing(p, before=18, after=10, line=1.25)
        size, color = 16, BLUE
    elif level == 2:
        set_paragraph_spacing(p, before=14, after=7, line=1.25)
        size, color = 13, BLUE
    else:
        set_paragraph_spacing(p, before=10, after=5, line=1.25)
        size, color = 12, DARK_BLUE
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=True)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, before=0, after=4, line=1.25)
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    run = p.add_run(text)
    set_run_font(run, size=11, color=BLACK)
    return p


def set_cell_text(cell, text, bold=False, color=BLACK):
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0, line=1.15)
    run = p.add_run(text)
    set_run_font(run, size=10, color=color, bold=bold)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def add_mode_table(doc):
    table = doc.add_table(rows=1, cols=3)
    set_table_borders(table)
    set_table_width(table, [1.1, 1.25, 4.15])
    headers = ["Mode", "Bits/symbol", "What to expect"]
    for i, header in enumerate(headers):
        shade_cell(table.rows[0].cells[i], TABLE_FILL)
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=DARK_BLUE)
    for mode, _, bits, desc in MODES:
        cells = table.add_row().cells
        set_cell_text(cells[0], mode, bold=True)
        set_cell_text(cells[1], bits)
        set_cell_text(cells[2], desc)
    add_para(doc, "", after=2)


def add_caption(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=2, after=8, line=1.15)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=9, color=MUTED, italic=True)


def add_image(doc, image_path, width=6.2):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=2, after=2, line=1.0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))


def add_page_break(doc):
    doc.add_page_break()


def set_section_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(footer, before=0, after=0, line=1.0)
    run = footer.add_run("ADALM-Pluto Dashboard Plot Guide")
    set_run_font(run, size=9, color=MUTED)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    set_section_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    set_paragraph_spacing(title, before=0, after=4, line=1.15)
    r = title.add_run("ADALM-Pluto Dashboard Plot Guide")
    set_run_font(r, size=24, color=BLACK, bold=True)

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, before=0, after=12, line=1.15)
    r = subtitle.add_run("ENGR 078 Final Project 1 live signal displays")
    set_run_font(r, size=12, color=MUTED)

    add_para(
        doc,
        "This guide explains the plots in the Pluto dashboard using only the class-connected signal ideas shown in the app: "
        "spectrum/PSD, I/Q constellation, eye diagram, SNR/noise, and throughput.",
    )
    add_para(
        doc,
        "Screenshots were captured from the running ADALM-Pluto dashboard in BPSK, QPSK, and 16-QAM modes. "
        "Because these are live Pluto captures, the exact point clouds and rates can change each time the app runs.",
    )

    add_heading(doc, "Quick Map", 1)
    add_mode_table(doc)
    add_heading(doc, "General Reading Order", 2)
    add_bullet(doc, "First check the status and modulation mode, so the expected number of symbol points makes sense.")
    add_bullet(doc, "Then check the spectrum to confirm signal energy is present near baseband.")
    add_bullet(doc, "Use the constellation and eye diagram to judge signal quality and timing.")
    add_bullet(doc, "Use the throughput counter to confirm that decoded bits are accumulating over time.")

    add_page_break(doc)

    fig_num = 1
    for plot_title, slug, meaning, bullets, takeaway in PLOTS:
        add_heading(doc, plot_title, 1)
        add_heading(doc, "What It Means", 2)
        add_para(doc, meaning)
        add_heading(doc, "How To Read It", 2)
        for item in bullets:
            add_bullet(doc, item)
        add_para(doc, f"Quick check: {takeaway}")

        for mode, safe, _, _ in MODES:
            image_path = IMG_DIR / f"{safe}_{slug}.png"
            add_image(doc, image_path)
            add_caption(doc, f"Figure {fig_num}. {plot_title} in {mode} mode.")
            fig_num += 1
        if plot_title != PLOTS[-1][0]:
            add_page_break(doc)

    add_page_break(doc)
    add_heading(doc, "Control Effects", 1)
    add_heading(doc, "Software AWGN SNR", 2)
    add_para(doc, "Lowering this slider adds more noise in software. In the spectrum, the floor rises. In the constellation, the green dots spread. In the eye diagram, the opening becomes less clean.")
    add_heading(doc, "Carrier Phase Offset", 2)
    add_para(doc, "Changing phase rotates the I/Q constellation. This is easiest to see on PSK modes because their information is carried strongly by phase.")
    add_heading(doc, "Carrier Frequency Offset", 2)
    add_para(doc, "Changing frequency offset shifts or smears the received signal over time. In the spectrum, the signal band moves away from center; in the constellation, points can rotate or blur.")
    add_heading(doc, "Throughput", 2)
    add_para(doc, "Throughput is a live rate counter, so it is expected to jump. The important sign is that the cumulative-bit curve keeps increasing while streaming.")

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
